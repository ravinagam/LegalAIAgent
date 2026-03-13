"""
Legal AI Agent — FastAPI Backend
Analyzes legal documents: summaries, clause extraction, risk analysis, AI chat.
"""
import hashlib
import json
import logging
import os
import traceback
import uuid
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import AsyncGenerator

import anthropic
from dotenv import load_dotenv
from fastapi import FastAPI, File, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, PlainTextResponse, StreamingResponse
from pydantic import BaseModel

load_dotenv()

logging.basicConfig(level=logging.INFO)
log = logging.getLogger("legal-ai")

# ---------------------------------------------------------------------------
# MongoDB (optional) — set MONGODB_URI env var to enable persistent storage
# ---------------------------------------------------------------------------
_MONGODB_URI = os.getenv("MONGODB_URI", "")
_mongo_db   = None   # set in startup
_grid_fs    = None   # AsyncIOMotorGridFSBucket for original file binaries

try:
    from motor.motor_asyncio import AsyncIOMotorClient, AsyncIOMotorGridFSBucket
    _MOTOR_OK = True
except ImportError:
    _MOTOR_OK = False
    if _MONGODB_URI:
        log.warning("motor not installed but MONGODB_URI is set — run: pip install motor")


async def _init_mongo() -> bool:
    global _mongo_db, _grid_fs
    if not _MOTOR_OK or not _MONGODB_URI:
        return False
    try:
        _client = AsyncIOMotorClient(_MONGODB_URI, serverSelectionTimeoutMS=5000)
        # derive db name from URI (last path segment before '?'), default "legalai"
        db_name = (_MONGODB_URI.split("/")[-1].split("?")[0]).strip() or "legalai"
        _mongo_db = _client[db_name]
        _grid_fs  = AsyncIOMotorGridFSBucket(_mongo_db, bucket_name="originals")
        await _client.admin.command("ping")
        log.info("MongoDB connected (db=%s)", db_name)
        return True
    except Exception as exc:
        log.error("MongoDB connection failed: %s — falling back to file storage", exc)
        _mongo_db = None
        _grid_fs  = None
        return False


async def _mongo_load_users() -> None:
    async for u in _mongo_db.users.find():
        u.pop("_id", None)
        users_db[u["username"]] = u


async def _mongo_save_user(username: str) -> None:
    data = {"_id": username, **users_db[username]}
    await _mongo_db.users.replace_one({"_id": username}, data, upsert=True)


async def _mongo_load_docs() -> None:
    async for d in _mongo_db.documents.find():
        d.pop("_id", None)
        documents[d["id"]] = d


async def _mongo_save_doc(doc: dict) -> None:
    data = {"_id": doc["id"], **doc}
    await _mongo_db.documents.replace_one({"_id": doc["id"]}, data, upsert=True)


async def _mongo_delete_doc(doc_id: str) -> None:
    await _mongo_db.documents.delete_one({"_id": doc_id})
    # delete original binary from GridFS
    cursor = _grid_fs.find({"metadata.doc_id": doc_id})
    async for gf in cursor:
        await _grid_fs.delete(gf._id)


async def _mongo_save_original(doc_id: str, ext: str, data: bytes) -> None:
    """Store (or replace) original binary in GridFS."""
    fname = f"{doc_id}_original{ext}"
    # delete any previous version
    cursor = _grid_fs.find({"metadata.doc_id": doc_id})
    async for gf in cursor:
        await _grid_fs.delete(gf._id)
    await _grid_fs.upload_from_stream(fname, data, metadata={"doc_id": doc_id})


async def _mongo_read_original(doc_id: str, ext: str) -> bytes | None:
    """Read original binary from GridFS; returns None if not found."""
    try:
        fname  = f"{doc_id}_original{ext}"
        stream = await _grid_fs.open_download_stream_by_name(fname)
        return await stream.read()
    except Exception:
        return None

# ---------------------------------------------------------------------------
# App & CORS
# ---------------------------------------------------------------------------
app = FastAPI(title="Legal AI Agent", version="1.0.0")


@app.on_event("startup")
async def startup_check():
    key = os.getenv("ANTHROPIC_API_KEY", "")
    if not key or not key.startswith("sk-"):
        log.error("=" * 60)
        log.error("ANTHROPIC_API_KEY is missing or invalid!")
        log.error("Add it to your .env file:  ANTHROPIC_API_KEY=sk-ant-...")
        log.error("=" * 60)
    else:
        log.info(f"Anthropic API key loaded (sk-...{key[-6:]})")

    if _MONGODB_URI:
        mongo_ok = await _init_mongo()
        if mongo_ok:
            users_db.clear()
            documents.clear()
            await _mongo_load_users()
            await _mongo_load_docs()
            log.info("Loaded %d users, %d documents from MongoDB",
                     len(users_db), len(documents))
        else:
            log.warning("MongoDB unavailable — using file storage fallback")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"],
)

# ---------------------------------------------------------------------------
# Storage root — use DATA_DIR env var if set (e.g. Render persistent disk)
# ---------------------------------------------------------------------------
_DATA_ROOT = Path(os.getenv("DATA_DIR", str(Path(__file__).parent)))
_DATA_ROOT.mkdir(parents=True, exist_ok=True)

# ---------------------------------------------------------------------------
# Auth — user store (persisted to users.json) + in-memory sessions
# ---------------------------------------------------------------------------
_USERS_FILE = _DATA_ROOT / "users.json"
users_db: dict[str, dict] = {}      # username → {username, password_hash, email, created_at}
sessions: dict[str, str]  = {}      # token → username


def _load_users() -> None:
    global users_db
    if _USERS_FILE.exists():
        try:
            users_db = json.loads(_USERS_FILE.read_text(encoding="utf-8"))
        except Exception:
            users_db = {}


def _save_users() -> None:
    _USERS_FILE.write_text(json.dumps(users_db, indent=2), encoding="utf-8")


def _hash_password(password: str) -> str:
    return hashlib.sha256(password.encode("utf-8")).hexdigest()


def _token_to_user(request: Request) -> str | None:
    auth = request.headers.get("Authorization", "")
    token = auth.removeprefix("Bearer ").strip()
    return sessions.get(token)


_load_users()

# ---------------------------------------------------------------------------
# In-memory document store: {doc_id: {...}}
# ---------------------------------------------------------------------------
documents: dict[str, dict] = {}

# ---------------------------------------------------------------------------
# Document disk persistence
# ---------------------------------------------------------------------------
_DOCS_DIR = _DATA_ROOT / "documents"
_DOCS_DIR.mkdir(exist_ok=True)


def _save_doc(doc: dict) -> None:
    try:
        (_DOCS_DIR / f"{doc['id']}.json").write_text(
            json.dumps(doc, indent=2), encoding="utf-8"
        )
    except Exception as exc:
        log.warning("Failed to save doc %s to disk: %s", doc.get("id"), exc)


def _delete_doc_file(doc_id: str) -> None:
    for name in [f"{doc_id}.json"] + [f"{doc_id}_original{e}" for e in _MIME_TYPES]:
        p = _DOCS_DIR / name
        try:
            if p.exists():
                p.unlink()
        except Exception as exc:
            log.warning("Failed to delete %s: %s", name, exc)


def _load_docs_from_disk() -> None:
    for p in sorted(_DOCS_DIR.glob("*.json")):
        try:
            doc = json.loads(p.read_text(encoding="utf-8"))
            documents[doc["id"]] = doc
            log.info("Loaded doc %s (%s) from disk", doc["id"], doc.get("filename", "?"))
        except Exception as exc:
            log.warning("Failed to load doc file %s: %s", p.name, exc)


_load_docs_from_disk()

# ---------------------------------------------------------------------------
# MIME types for original file download
# ---------------------------------------------------------------------------
_MIME_TYPES: dict[str, str] = {
    ".pdf":  "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc":  "application/msword",
    ".txt":  "text/plain; charset=utf-8",
    ".md":   "text/markdown; charset=utf-8",
}

# ---------------------------------------------------------------------------
# PDF report generation (reportlab)
# ---------------------------------------------------------------------------
try:
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        HRFlowable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
    )
    _REPORTLAB_OK = True
except ImportError:
    _REPORTLAB_OK = False
    log.warning("reportlab not installed — PDF reports unavailable. Run: pip install reportlab")


def _safe(text) -> str:
    """Encode text for reportlab: Latin-1 safe + XML-escaped for Paragraph content."""
    s = str(text or "").encode("latin-1", errors="replace").decode("latin-1")
    s = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return s


def _generate_pdf_report(doc: dict) -> bytes:
    from io import BytesIO
    a   = doc.get("analysis", {})
    buf = BytesIO()

    pdf = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
        title=f"Analysis — {doc.get('filename', '')}",
        author="Legal AI Agent",
    )

    styles = getSampleStyleSheet()
    BLUE   = rl_colors.HexColor("#3b82f6")
    DARK   = rl_colors.HexColor("#1e293b")
    MUTED  = rl_colors.HexColor("#64748b")
    BORDER = rl_colors.HexColor("#e2e8f0")
    BG     = rl_colors.HexColor("#f8fafc")
    RED    = rl_colors.HexColor("#ef4444")
    AMBER  = rl_colors.HexColor("#f59e0b")
    GREEN  = rl_colors.HexColor("#10b981")
    RISK_C    = {"High": RED,      "Medium": AMBER,    "Low": GREEN}
    RISK_HEX  = {"High": "ef4444", "Medium": "f59e0b", "Low": "10b981"}
    MUTED_HEX = "64748b"

    T  = lambda s, **kw: ParagraphStyle(s, parent=styles["Normal"], **kw)
    h1 = T("H1", fontSize=13, fontName="Helvetica-Bold", textColor=DARK,
           spaceBefore=14, spaceAfter=6)
    lbl = T("Lbl", fontSize=9,  fontName="Helvetica-Bold", textColor=MUTED)
    val = T("Val", fontSize=10, textColor=DARK)
    body= T("Body", fontSize=10, textColor=DARK, leading=15, spaceAfter=4)
    sm  = T("Sm", fontSize=9,  textColor=MUTED,  leading=13)

    def hr(thick=0.5): return HRFlowable(width="100%", thickness=thick, color=BORDER, spaceAfter=8)

    story = []

    # ── Title block ──────────────────────────────────────────────────────────
    story.append(Paragraph("Legal Document Analysis Report",
        T("Title", fontSize=20, fontName="Helvetica-Bold", textColor=DARK,
          alignment=TA_CENTER, spaceAfter=4)))
    story.append(Paragraph("Generated by Legal AI Agent &bull; Powered by Claude",
        T("Sub", fontSize=10, textColor=MUTED, alignment=TA_CENTER, spaceAfter=16)))
    story.append(HRFlowable(width="100%", thickness=2, color=BLUE, spaceAfter=18))

    # ── Document metadata table ───────────────────────────────────────────────
    risk   = a.get("risk_level", "Unknown")
    rc_hex = RISK_HEX.get(risk, MUTED_HEX)
    parties_str = _safe(", ".join(a.get("parties", [])) or "—")
    dt_str = ""
    if doc.get("uploaded_at"):
        try:
            dt_str = datetime.fromisoformat(doc["uploaded_at"]).strftime("%d %b %Y %H:%M")
        except Exception:
            dt_str = doc["uploaded_at"]

    rows = [
        (_safe("Document"),     _safe(doc.get("filename", ""))),
        (_safe("Type"),         _safe(a.get("document_type") or "—")),
        (_safe("Parties"),      parties_str),
        (_safe("Effective"),    _safe(a.get("effective_date") or "—")),
        (_safe("Expiry"),       _safe(a.get("expiry_date")    or "—")),
        (_safe("Governing Law"),_safe(a.get("governing_law")  or "—")),
        (_safe("Risk Level"),   f'<font color="#{rc_hex}">{_safe(risk)}</font>'),
    ]
    if dt_str:
        rows.append((_safe("Analyzed"), _safe(dt_str)))

    tdata = [[Paragraph(r, lbl), Paragraph(v, val)] for r, v in rows]
    t = Table(tdata, colWidths=[3.8*cm, None])
    t.setStyle(TableStyle([
        ("ROWBACKGROUNDS", (0,0), (-1,-1), [BG, rl_colors.white]),
        ("BOX",            (0,0), (-1,-1), 0.5, BORDER),
        ("TOPPADDING",     (0,0), (-1,-1), 6),
        ("BOTTOMPADDING",  (0,0), (-1,-1), 6),
        ("LEFTPADDING",    (0,0), (-1,-1), 10),
        ("RIGHTPADDING",   (0,0), (-1,-1), 10),
    ]))
    story += [t, Spacer(1, 18)]

    # ── Summary ──────────────────────────────────────────────────────────────
    if a.get("summary"):
        story += [Paragraph("Executive Summary", h1), hr(),
                  Paragraph(_safe(a["summary"]), body), Spacer(1, 10)]

    # ── Key obligations ───────────────────────────────────────────────────────
    if a.get("key_obligations"):
        story += [Paragraph("Key Obligations", h1), hr()]
        odata = [[Paragraph("Party", lbl), Paragraph("Obligation", lbl)]]
        for o in a["key_obligations"]:
            odata.append([Paragraph(_safe(o.get("party","")), val),
                          Paragraph(_safe(o.get("obligation","")), body)])
        ot = Table(odata, colWidths=[4*cm, None])
        ot.setStyle(TableStyle([
            ("BACKGROUND",   (0,0), (-1,0), BLUE),
            ("TEXTCOLOR",    (0,0), (-1,0), rl_colors.white),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[BG, rl_colors.white]),
            ("BOX",          (0,0), (-1,-1), 0.5, BORDER),
            ("LINEBELOW",    (0,0), (-1,0),  1,   rl_colors.HexColor("#2563eb")),
            ("TOPPADDING",   (0,0), (-1,-1), 6),
            ("BOTTOMPADDING",(0,0), (-1,-1), 6),
            ("LEFTPADDING",  (0,0), (-1,-1), 8),
        ]))
        story += [ot, Spacer(1, 10)]

    # ── Risk factors ──────────────────────────────────────────────────────────
    if a.get("risk_factors"):
        story += [Paragraph("Risk Factors", h1), hr()]
        for rf in a["risk_factors"]:
            story.append(Paragraph(f"&bull; {_safe(rf)}", body))
        story.append(Spacer(1, 10))

    # ── Clauses ───────────────────────────────────────────────────────────────
    if a.get("clauses"):
        story += [Paragraph("Clause Analysis", h1), hr()]
        for c in a["clauses"]:
            rl_val = c.get("risk_level", "Low")
            rc2    = RISK_C.get(rl_val, MUTED)
            rc2_hex = RISK_HEX.get(rl_val, MUTED_HEX)
            story.append(Paragraph(
                f'{_safe(c.get("type",""))} &mdash; {_safe(c.get("title",""))}'
                f'<font color="#{rc2_hex}"> [{_safe(rl_val)} Risk]</font>',
                T("CT", fontSize=10, fontName="Helvetica-Bold", textColor=DARK,
                  spaceBefore=8, spaceAfter=3)))
            story.append(Paragraph(_safe(c.get("content","")),
                T("CC", fontSize=9, textColor=MUTED, leading=13, leftIndent=10)))
            if c.get("risk_note") and rl_val != "Low":
                story.append(Paragraph(f"&bull; {_safe(c['risk_note'])}",
                    T("RN", fontSize=9, textColor=rc2, fontName="Helvetica-Oblique",
                      leftIndent=10, spaceAfter=4)))
            story.append(hr(0.3))

    pdf.build(story)
    return buf.getvalue()

# ---------------------------------------------------------------------------
# Anthropic client  (async — required for use inside FastAPI async endpoints)
# ---------------------------------------------------------------------------
client = anthropic.AsyncAnthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
MODEL = "claude-sonnet-4-6"

# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------

# Validate required packages are available at startup (not silently at upload time)
try:
    from pypdf import PdfReader as _PdfReader
    _PYPDF_OK = True
except ImportError:
    _PYPDF_OK = False
    log.warning("pypdf not installed — PDF uploads will fail. Run: pip install pypdf")

try:
    from docx import Document as _DocxDocument
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False
    log.warning("python-docx not installed — DOCX uploads will fail. Run: pip install python-docx")


def _extract_pdf(data: bytes) -> str:
    if not _PYPDF_OK:
        raise HTTPException(
            status_code=500,
            detail="pypdf is not installed on the server. Run: pip install pypdf",
        )
    try:
        reader = _PdfReader(BytesIO(data))
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Cannot open PDF: {exc}")

    pages_text: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            t = page.extract_text() or ""
            if t.strip():
                pages_text.append(t)
        except Exception as exc:
            log.warning("PDF page %d extraction skipped: %s", i, exc)

    result = "\n\n".join(pages_text).strip()
    log.info("PDF: extracted %d chars from %d/%d pages", len(result), len(pages_text), len(reader.pages))

    if not result:
        raise HTTPException(
            status_code=422,
            detail=(
                "No text could be extracted from this PDF. "
                "It may be a scanned image-only PDF. "
                "Try converting it to a text-based PDF or copy-paste the text into a .txt file."
            ),
        )
    return result


def _create_docx_from_text(text: str) -> bytes:
    """Create a simple DOCX file from plain text (one paragraph per non-empty line)."""
    doc = _DocxDocument()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _extract_docx(data: bytes) -> str:
    if not _DOCX_OK:
        raise HTTPException(
            status_code=500,
            detail="python-docx is not installed on the server. Run: pip install python-docx",
        )
    try:
        doc = _DocxDocument(BytesIO(data))
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Cannot open DOCX: {exc}")

    parts: list[str] = []

    # Top-level paragraphs
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)

    # Tables — iterate every cell in every table
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text.strip()
                    if t:
                        parts.append(t)

    # Headers and footers
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf:
                for p in hf.paragraphs:
                    t = p.text.strip()
                    if t:
                        parts.append(t)

    result = "\n\n".join(parts).strip()
    log.info("DOCX: extracted %d chars from %d text blocks", len(result), len(parts))

    if not result:
        raise HTTPException(
            status_code=422,
            detail=(
                "No text could be extracted from this DOCX. "
                "The document may be empty or contain only images/embedded objects."
            ),
        )
    return result


def extract_text(data: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(data)
    if ext in (".docx", ".doc"):
        return _extract_docx(data)
    if ext in (".txt", ".md"):
        text = data.decode("utf-8", errors="replace").strip()
        if not text:
            raise HTTPException(status_code=422, detail="The text file is empty.")
        return text
    raise HTTPException(
        status_code=415,
        detail=f"Unsupported file type '{ext}'. Supported: PDF, DOCX, DOC, TXT, MD.",
    )


# ---------------------------------------------------------------------------
# Claude analysis helpers
# ---------------------------------------------------------------------------

ANALYSIS_SYSTEM = """You are an expert legal analyst AI. Your task is to analyze legal documents
and produce structured JSON output. Always respond with ONLY valid JSON — no markdown fences,
no extra text. Be thorough, precise, and highlight risks clearly."""

# NOTE: Use plain string concatenation for the document text — never .format() or
# f-strings — because legal documents routinely contain { } characters (e.g. ${1,000},
# Section {a}(i), exhibit labels) which would crash Python's str.format().
ANALYSIS_PROMPT_PREFIX = """Analyze the following legal document and return a JSON object with
exactly this structure (no extra keys):

{
  "document_type": "string (e.g. Service Agreement, NDA, Employment Contract, Lease, etc.)",
  "parties": ["list of party names as strings"],
  "effective_date": "string or null",
  "expiry_date": "string or null",
  "governing_law": "string or null",
  "summary": "3-5 sentence executive summary of the document",
  "risk_level": "Low | Medium | High",
  "risk_factors": ["list of specific risk concerns as strings"],
  "key_obligations": [
    {"party": "party name", "obligation": "what they must do"}
  ],
  "clauses": [
    {
      "type": "clause category (e.g. Termination, Payment, Confidentiality, Liability, IP, Non-Compete, Dispute Resolution, Force Majeure, Indemnification, Amendment)",
      "title": "short clause title",
      "content": "verbatim or close-paraphrase of the clause text",
      "risk_level": "Low | Medium | High",
      "risk_note": "brief note on why this clause is risky or what to watch out for (null if low risk)"
    }
  ]
}

Document to analyze:
---
"""

ANALYSIS_PROMPT_SUFFIX = "\n---"


async def analyze_document(text: str, filename: str) -> dict:
    """Call Claude to analyze a legal document. Returns structured dict."""
    truncated = text[:150_000]
    # Build prompt by concatenation — NOT .format() — to avoid KeyError on { } in docs
    prompt = ANALYSIS_PROMPT_PREFIX + truncated + ANALYSIS_PROMPT_SUFFIX

    log.info("Sending %d chars to Claude for analysis…", len(truncated))
    response = await client.messages.create(
        model=MODEL,
        max_tokens=16000,
        system=ANALYSIS_SYSTEM,
        messages=[{"role": "user", "content": prompt}],
    )
    log.info("Claude response received (stop_reason=%s)", response.stop_reason)

    raw = next(
        (b.text for b in response.content if b.type == "text"), "{}"
    )

    # Strip accidental markdown fences if model added them
    raw = raw.strip()
    if raw.startswith("```"):
        raw = raw.split("```", 2)[1]
        if raw.startswith("json"):
            raw = raw[4:]
        raw = raw.rsplit("```", 1)[0].strip()

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        # Fallback minimal structure
        return {
            "document_type": "Unknown",
            "parties": [],
            "effective_date": None,
            "expiry_date": None,
            "governing_law": None,
            "summary": "Analysis could not be parsed. Please review the document manually.",
            "risk_level": "Medium",
            "risk_factors": [],
            "key_obligations": [],
            "clauses": [],
        }


async def stream_chat_response(
    doc_id: str, user_message: str
) -> AsyncGenerator[str, None]:
    """Stream a Claude response for document chat as SSE events."""
    doc = documents.get(doc_id)
    if not doc:
        yield f"data: {json.dumps({'error': 'Document not found'})}\n\n"
        return

    analysis = doc.get("analysis", {})
    doc_type = analysis.get("document_type", "Unknown")
    parties  = ", ".join(analysis.get("parties", []))

    # Build system prompt by concatenation — NOT .format() — to avoid KeyError on { } in docs
    system_prompt = (
        "You are an expert AI legal assistant. You are helping a user understand a legal document. "
        "Be clear, precise, and helpful. When referencing specific clauses, cite them. "
        "Flag any risks or obligations the user should be aware of.\n\n"
        f"Document type: {doc_type}\n"
        f"Parties: {parties}\n\n"
        "Full document text:\n---\n"
        + doc["text"][:120_000]
        + "\n---"
    )

    # Build message history (keep last 20 turns)
    history = doc.get("chat_history", [])
    messages = history[-20:] + [{"role": "user", "content": user_message}]

    # Stream via Claude
    async with client.messages.stream(
        model=MODEL,
        max_tokens=4096,
        system=system_prompt,
        messages=messages,
    ) as stream:
        full_response = ""
        async for text_chunk in stream.text_stream:
            full_response += text_chunk
            yield f"data: {json.dumps({'delta': text_chunk})}\n\n"

    # Persist conversation history
    doc["chat_history"].append({"role": "user", "content": user_message})
    doc["chat_history"].append({"role": "assistant", "content": full_response})
    _save_doc(doc)

    yield f"data: {json.dumps({'done': True})}\n\n"


# ---------------------------------------------------------------------------
# Request / Response models
# ---------------------------------------------------------------------------

class RegisterRequest(BaseModel):
    username: str
    password: str
    email: str = ""


class LoginRequest(BaseModel):
    username: str
    password: str


class UpdateDocRequest(BaseModel):
    text: str
    reanalyze: bool = False


class ChatRequest(BaseModel):
    message: str


class ClearHistoryRequest(BaseModel):
    pass


# ---------------------------------------------------------------------------
# API Routes
# ---------------------------------------------------------------------------

@app.post("/api/register")
async def register(body: RegisterRequest):
    username = body.username.strip().lower()
    if not username or not body.password:
        raise HTTPException(status_code=400, detail="Username and password are required.")
    if len(username) < 3:
        raise HTTPException(status_code=400, detail="Username must be at least 3 characters.")
    if len(body.password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters.")
    if username in users_db:
        raise HTTPException(status_code=409, detail="Username already exists.")
    users_db[username] = {
        "username": username,
        "password_hash": _hash_password(body.password),
        "email": body.email.strip(),
        "created_at": datetime.now().isoformat(),
    }
    if _mongo_db is not None:
        await _mongo_save_user(username)
    else:
        _save_users()
    token = str(uuid.uuid4())
    sessions[token] = username
    log.info("New user registered: %s", username)
    return {"token": token, "username": username}


@app.post("/api/login")
async def login(body: LoginRequest):
    username = body.username.strip().lower()
    user = users_db.get(username)
    if not user or user["password_hash"] != _hash_password(body.password):
        raise HTTPException(status_code=401, detail="Invalid username or password.")
    token = str(uuid.uuid4())
    sessions[token] = username
    log.info("User logged in: %s", username)
    return {"token": token, "username": username}


@app.post("/api/logout")
async def logout(request: Request):
    auth = request.headers.get("Authorization", "")
    token = auth.removeprefix("Bearer ").strip()
    sessions.pop(token, None)
    return {"status": "logged out"}


@app.get("/api/me")
async def me(request: Request):
    username = _token_to_user(request)
    if not username:
        raise HTTPException(status_code=401, detail="Not authenticated.")
    user = users_db.get(username, {})
    return {"username": username, "email": user.get("email", "")}


@app.post("/api/upload")
async def upload_document(request: Request, file: UploadFile = File(...)):
    """Upload a legal document, extract text, and auto-analyze with Claude."""
    try:
        data = await file.read()
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Failed to read uploaded file: {exc}")

    if len(data) > 20 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="File too large (max 20 MB).")

    # ── Text extraction ──────────────────────────────────────────────────
    try:
        text = extract_text(data, file.filename or "document.txt")
    except HTTPException:
        raise
    except Exception as exc:
        log.error("Text extraction error: %s", traceback.format_exc())
        raise HTTPException(status_code=422, detail=f"Text extraction failed: {exc}")

    log.info("Extracted %d chars from '%s'", len(text), file.filename)

    # ── Claude analysis ──────────────────────────────────────────────────
    try:
        analysis = await analyze_document(text, file.filename or "document")
    except anthropic.AuthenticationError:
        log.error("Anthropic AuthenticationError — check ANTHROPIC_API_KEY in .env")
        raise HTTPException(
            status_code=401,
            detail="Invalid Anthropic API key. Open .env and set ANTHROPIC_API_KEY=sk-ant-...",
        )
    except anthropic.RateLimitError as exc:
        log.error("Anthropic RateLimitError: %s", exc)
        raise HTTPException(status_code=429, detail="Anthropic rate limit reached. Wait a moment and retry.")
    except anthropic.APIConnectionError as exc:
        log.error("Anthropic APIConnectionError: %s", exc)
        raise HTTPException(status_code=503, detail=f"Cannot reach Anthropic API: {exc}")
    except anthropic.BadRequestError as exc:
        log.error("Anthropic BadRequestError: %s\n%s", exc, traceback.format_exc())
        raise HTTPException(status_code=400, detail=f"Anthropic rejected the request: {exc}")
    except anthropic.APIStatusError as exc:
        log.error("Anthropic APIStatusError %s: %s", exc.status_code, exc)
        raise HTTPException(status_code=502, detail=f"Anthropic API error {exc.status_code}: {exc}")
    except Exception as exc:
        log.error("Unexpected analysis error:\n%s", traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Analysis failed: {exc}")

    doc_id = str(uuid.uuid4())
    owner = _token_to_user(request) or ""
    uploaded_at = datetime.now().isoformat()
    original_ext = Path(file.filename or "document.txt").suffix.lower()

    # Save original binary (GridFS if MongoDB, else disk)
    try:
        if _mongo_db is not None:
            await _mongo_save_original(doc_id, original_ext, data)
        else:
            (_DOCS_DIR / f"{doc_id}_original{original_ext}").write_bytes(data)
    except Exception as exc:
        log.warning("Could not save original file for %s: %s", doc_id, exc)

    documents[doc_id] = {
        "id": doc_id,
        "filename": file.filename,
        "size": len(data),
        "text": text,
        "analysis": analysis,
        "chat_history": [],
        "owner": owner,
        "uploaded_at": uploaded_at,
        "original_ext": original_ext,
    }
    if _mongo_db is not None:
        await _mongo_save_doc(documents[doc_id])
    else:
        _save_doc(documents[doc_id])

    log.info("Document %s stored successfully (owner=%s)", doc_id, owner or "anonymous")
    return {
        "id": doc_id,
        "filename": file.filename,
        "size": len(data),
        "analysis": analysis,
        "uploaded_at": uploaded_at,
    }


@app.get("/api/documents")
async def list_documents(request: Request):
    """Return metadata for documents owned by the authenticated user."""
    owner = _token_to_user(request) or ""
    return [
        {
            "id": d["id"],
            "filename": d["filename"],
            "size": d["size"],
            "document_type": d["analysis"].get("document_type", "Unknown"),
            "risk_level": d["analysis"].get("risk_level", "Unknown"),
            "parties": d["analysis"].get("parties", []),
            "uploaded_at": d.get("uploaded_at"),
        }
        for d in documents.values()
        if not owner or d.get("owner") == owner
    ]


@app.get("/api/documents/{doc_id}")
async def get_document(doc_id: str):
    """Return full analysis for a document."""
    doc = documents.get(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found.")
    return {
        "id": doc["id"],
        "filename": doc["filename"],
        "size": doc["size"],
        "analysis": doc["analysis"],
        "chat_history": doc["chat_history"],
        "uploaded_at": doc.get("uploaded_at"),
    }


@app.get("/api/documents/{doc_id}/download")
async def download_original_file(doc_id: str):
    """Download the original uploaded file (PDF, DOCX, TXT, etc.)."""
    doc = documents.get(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found.")

    ext  = doc.get("original_ext", Path(doc.get("filename", "document.txt")).suffix.lower())
    mime = _MIME_TYPES.get(ext, "application/octet-stream")
    from fastapi.responses import Response as _Resp

    # Try GridFS first, then disk
    binary = None
    if _mongo_db is not None:
        binary = await _mongo_read_original(doc_id, ext)
    else:
        p = _DOCS_DIR / f"{doc_id}_original{ext}"
        if p.exists():
            binary = p.read_bytes()

    if binary is not None:
        return _Resp(
            content=binary,
            media_type=mime,
            headers={"Content-Disposition": f'attachment; filename="{doc["filename"]}"'},
        )

    # Fallback — no binary stored: serve extracted text as .txt
    stem = Path(doc["filename"]).stem
    safe_stem = "".join(c if c.isalnum() or c in "-_ " else "_" for c in stem).strip() or "document"
    return PlainTextResponse(
        content=doc["text"],
        headers={"Content-Disposition": f'attachment; filename="{safe_stem}.txt"'},
    )


@app.get("/api/documents/{doc_id}/report")
async def download_pdf_report(doc_id: str):
    """Generate and download a PDF analysis report for a document."""
    if not _REPORTLAB_OK:
        raise HTTPException(
            status_code=501,
            detail="PDF generation unavailable. Install reportlab: pip install reportlab",
        )
    doc = documents.get(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found.")
    try:
        pdf_bytes = _generate_pdf_report(doc)
    except Exception as exc:
        log.error("PDF report generation failed: %s", exc)
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {exc}")

    stem = Path(doc["filename"]).stem
    safe = "".join(c if c.isalnum() or c in "-_ " else "_" for c in stem).strip() or "document"
    from fastapi.responses import Response
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{safe}_analysis.pdf"'},
    )


@app.get("/api/documents/{doc_id}/text")
async def get_document_text(doc_id: str):
    """Return the raw extracted text for a document (for editing)."""
    doc = documents.get(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found.")
    return {"text": doc["text"]}


@app.put("/api/documents/{doc_id}")
async def update_document(doc_id: str, body: UpdateDocRequest):
    """Update document text and optionally re-analyze with Claude."""
    doc = documents.get(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found.")

    if not body.text.strip():
        raise HTTPException(status_code=400, detail="Document text cannot be empty.")

    doc["text"] = body.text
    doc["size"] = len(body.text.encode("utf-8"))

    if body.reanalyze:
        try:
            analysis = await analyze_document(body.text, doc["filename"])
        except anthropic.AuthenticationError:
            raise HTTPException(status_code=401, detail="Invalid Anthropic API key.")
        except anthropic.RateLimitError:
            raise HTTPException(status_code=429, detail="Anthropic rate limit reached. Wait a moment and retry.")
        except anthropic.APIConnectionError as exc:
            raise HTTPException(status_code=503, detail=f"Cannot reach Anthropic API: {exc}")
        except Exception as exc:
            log.error("Re-analysis error:\n%s", traceback.format_exc())
            raise HTTPException(status_code=500, detail=f"Re-analysis failed: {exc}")
        doc["analysis"] = analysis
        doc["chat_history"] = []   # clear chat — document content changed

    # Regenerate the downloadable binary so it reflects the edited text
    ext = doc.get("original_ext", Path(doc.get("filename", "document.txt")).suffix.lower())
    try:
        if ext in (".docx", ".doc"):
            new_binary = _create_docx_from_text(body.text)
        elif ext in (".txt", ".md"):
            new_binary = body.text.encode("utf-8")
        else:
            new_binary = None  # PDF — not regenerated

        if new_binary is not None:
            if _mongo_db is not None:
                await _mongo_save_original(doc_id, ext, new_binary)
            else:
                (_DOCS_DIR / f"{doc_id}_original{ext}").write_bytes(new_binary)
    except Exception as exc:
        log.warning("Could not update original file for %s: %s", doc_id, exc)

    if _mongo_db is not None:
        await _mongo_save_doc(doc)
    else:
        _save_doc(doc)
    log.info("Document %s updated (reanalyze=%s)", doc_id, body.reanalyze)
    return {
        "id": doc["id"],
        "filename": doc["filename"],
        "size": doc["size"],
        "analysis": doc["analysis"],
        "chat_history": doc["chat_history"],
        "uploaded_at": doc.get("uploaded_at"),
    }


@app.delete("/api/documents/{doc_id}")
async def delete_document(doc_id: str):
    """Remove a document from memory."""
    if doc_id not in documents:
        raise HTTPException(status_code=404, detail="Document not found.")
    del documents[doc_id]
    if _mongo_db is not None:
        await _mongo_delete_doc(doc_id)
    else:
        _delete_doc_file(doc_id)
    return {"status": "deleted"}


@app.post("/api/documents/{doc_id}/chat")
async def chat_with_document(doc_id: str, body: ChatRequest):
    """Stream an AI response about the document as Server-Sent Events."""
    if doc_id not in documents:
        raise HTTPException(status_code=404, detail="Document not found.")
    return StreamingResponse(
        stream_chat_response(doc_id, body.message),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


@app.delete("/api/documents/{doc_id}/chat")
async def clear_chat_history(doc_id: str):
    """Clear the conversation history for a document."""
    doc = documents.get(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found.")
    doc["chat_history"] = []
    if _mongo_db is not None:
        await _mongo_save_doc(doc)
    else:
        _save_doc(doc)
    return {"status": "cleared"}


@app.get("/api/health")
async def health():
    return {"status": "ok", "model": MODEL}


@app.get("/api/test")
async def test_claude():
    """Quick smoke-test: sends a tiny message to Claude and returns the result."""
    try:
        response = await client.messages.create(
            model=MODEL,
            max_tokens=64,
            messages=[{"role": "user", "content": "Reply with just the word: OK"}],
        )
        reply = next((b.text for b in response.content if b.type == "text"), "")
        return {"status": "ok", "model": MODEL, "reply": reply}
    except anthropic.AuthenticationError:
        raise HTTPException(status_code=401, detail="Invalid ANTHROPIC_API_KEY")
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Claude test failed: {exc}")


# ---------------------------------------------------------------------------
# Serve frontend
# ---------------------------------------------------------------------------

@app.get("/", response_class=HTMLResponse)
async def serve_frontend():
    index_path = Path(__file__).parent / "index.html"
    if index_path.exists():
        return HTMLResponse(
            content=index_path.read_text(encoding="utf-8"),
            headers={"Cache-Control": "no-store, no-cache, must-revalidate"},
        )
    return HTMLResponse("<h1>index.html not found</h1>", status_code=404)
